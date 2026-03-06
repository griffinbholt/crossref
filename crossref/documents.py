import csv
import json
import logging
import os
import re

from dataclasses import dataclass
from typing import Callable

logger = logging.getLogger(__name__)


@dataclass
class Passage:
    text: str
    index: int
    label: str | None = None


class Document:
    """A document represented as an ordered list of passages with optional structural labels."""

    def __init__(self, passages: list[Passage], title: str = ""):
        self.passages = passages
        self.title = title

    def __repr__(self) -> str:
        return f"Document(title={self.title!r}, passages={len(self.passages)})"

    def __len__(self) -> int:
        return len(self.passages)

    def __getitem__(self, idx: int) -> Passage:
        return self.passages[idx]

    def __iter__(self):
        return iter(self.passages)

    def texts(self) -> list[str]:
        return [p.text for p in self.passages]

    def labels(self) -> list[str | None]:
        return [p.label for p in self.passages]

    # ------------------------------------------------------------------
    # Public constructors
    # ------------------------------------------------------------------

    @classmethod
    def from_file(
        cls,
        path: str,
        splitter: str | Callable = "line",
        title: str | None = None,
    ) -> 'Document':
        """Load a document from a file or directory.

        Format is auto-detected from the file extension:
          .md               Markdown — heading hierarchy becomes passage labels.
          .pdf              PDF — page number becomes the label. Requires pypdf.
          .docx             Word — heading hierarchy becomes labels. Requires python-docx.
          .html / .htm      HTML — heading hierarchy becomes labels. Requires beautifulsoup4.
          .epub             EPUB — heading hierarchy becomes labels.
                            Requires ebooklib and beautifulsoup4.
          (directory)       Recursively loads all supported files.
          (anything else)   Treated as plain text.

        For CSV, TSV, JSON, and JSONL files use the dedicated constructors
        from_csv(), from_json(), and from_jsonl() instead, which require
        the schema parameters (text_column / text_key) as explicit arguments.

        Args:
            path: Path to a file or directory.
            splitter: How to split content into passages. One of 'line', 'paragraph',
                      'section', 'sentence', or a callable that takes a string and
                      returns a list of strings.
            title: Document title. Inferred from filename/dirname if not provided.
        """
        path = os.path.expanduser(path)
        if os.path.isdir(path):
            return cls._from_directory(path, splitter, title)
        ext = os.path.splitext(path)[1].lower()
        if ext in ('.csv', '.tsv', '.json', '.jsonl', '.ndjson'):
            raise ValueError(
                f"Use Document.from_csv() or Document.from_json()/from_jsonl() "
                f"to load {ext!r} files — they require schema parameters "
                f"(text_column or text_key) that must be provided explicitly."
            )
        if ext == '.md':
            return cls._from_markdown(path, splitter, title)
        if ext == '.pdf':
            return cls._from_pdf(path, splitter, title)
        if ext == '.docx':
            return cls._from_docx(path, splitter, title)
        if ext in ('.html', '.htm'):
            return cls._from_html(path, splitter, title)
        if ext == '.epub':
            return cls._from_epub(path, splitter, title)
        return cls._from_text(path, splitter, title)

    @classmethod
    def from_passages(
        cls,
        texts: list[str],
        labels: list[str] | None = None,
        title: str = "",
    ) -> 'Document':
        """Construct a Document directly from pre-split passage texts.

        Args:
            texts: List of passage strings.
            labels: Optional structural labels, one per passage.
            title: Document title.
        """
        if labels is not None and len(labels) != len(texts):
            raise ValueError("labels must have the same length as texts")
        passages = [
            Passage(text=t, index=i, label=labels[i] if labels else None)
            for i, t in enumerate(texts)
        ]
        return cls(passages, title=title)

    @classmethod
    def from_csv(
        cls,
        path: str,
        text_column: str,
        label_column: str | None = None,
        delimiter: str | None = None,
        title: str | None = None,
    ) -> 'Document':
        """Load a document from a CSV or TSV file. Each row becomes one passage.

        Args:
            path: Path to the file. Tab delimiter is inferred automatically for .tsv.
            text_column: Name of the column containing passage text (required).
            label_column: Name of the column to use as passage label (optional).
            delimiter: Field delimiter. Defaults to tab for .tsv, comma for all others.
            title: Document title. Inferred from filename if not provided.
        """
        path = os.path.expanduser(path)
        ext = os.path.splitext(path)[1].lower()
        _delim = delimiter or ('\t' if ext == '.tsv' else ',')
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = []
        with open(path, 'r', newline='', encoding='utf-8') as f:
            for row in csv.DictReader(f, delimiter=_delim):
                text = (row.get(text_column) or '').strip()
                if not text:
                    continue
                label = (row.get(label_column) or '').strip() or None if label_column else None
                passages.append(Passage(text=text, index=len(passages), label=label))
        return cls(passages, title=doc_title)

    @classmethod
    def from_json(
        cls,
        path: str,
        text_key: str,
        label_key: str | None = None,
        title: str | None = None,
    ) -> 'Document':
        """Load a document from a JSON file. Each item in the array becomes one passage.

        The file may be a JSON array directly, or a dict wrapping an array
        (e.g. {"verses": [...]}), in which case the first list value is used.

        Args:
            path: Path to the .json file.
            text_key: Key whose value is the passage text (required).
            label_key: Key whose value is the passage label (optional).
            title: Document title. Inferred from filename if not provided.
        """
        path = os.path.expanduser(path)
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, dict):
            data = next((v for v in data.values() if isinstance(v, list)), [data])
        passages = []
        for item in data:
            text = str(item.get(text_key) or '').strip()
            if not text:
                continue
            label = str(item.get(label_key) or '').strip() or None if label_key else None
            passages.append(Passage(text=text, index=len(passages), label=label))
        return cls(passages, title=doc_title)

    @classmethod
    def from_jsonl(
        cls,
        path: str,
        text_key: str,
        label_key: str | None = None,
        title: str | None = None,
    ) -> 'Document':
        """Load a document from a JSONL (newline-delimited JSON) file.

        Each non-empty line must be a JSON object. Each line becomes one passage.

        Args:
            path: Path to the .jsonl or .ndjson file.
            text_key: Key whose value is the passage text (required).
            label_key: Key whose value is the passage label (optional).
            title: Document title. Inferred from filename if not provided.
        """
        path = os.path.expanduser(path)
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = []
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                item = json.loads(line)
                text = str(item.get(text_key) or '').strip()
                if not text:
                    continue
                label = str(item.get(label_key) or '').strip() or None if label_key else None
                passages.append(Passage(text=text, index=len(passages), label=label))
        return cls(passages, title=doc_title)

    # ------------------------------------------------------------------
    # Format-specific internal loaders
    # ------------------------------------------------------------------

    @classmethod
    def _from_markdown(cls, path, splitter, title) -> 'Document':
        logger.debug("Parsing markdown: %s", path)
        with open(path, 'r') as f:
            lines = [line.rstrip('\n') for line in f]
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = _parse_markdown(lines, splitter)
        for i, p in enumerate(passages):
            p.index = i
        doc = cls(passages, title=doc_title)
        logger.info("Loaded '%s': %d passages", doc_title, len(passages))
        return doc

    @classmethod
    def _from_text(cls, path, splitter, title) -> 'Document':
        logger.debug("Parsing text: %s", path)
        with open(path, 'r') as f:
            content = f.read()
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = []
        for t in _apply_splitter(content.splitlines(), splitter):
            if t.strip():
                passages.append(Passage(text=t, index=len(passages)))
        doc = cls(passages, title=doc_title)
        logger.info("Loaded '%s': %d passages", doc_title, len(passages))
        return doc

    @classmethod
    def _from_pdf(cls, path, splitter, title) -> 'Document':
        """Requires: pypdf"""
        from pypdf import PdfReader
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = _parse_pdf(PdfReader(path), splitter)
        for i, p in enumerate(passages):
            p.index = i
        return cls(passages, title=doc_title)

    @classmethod
    def _from_docx(cls, path, splitter, title) -> 'Document':
        """Requires: python-docx"""
        from docx import Document as DocxDocument
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        passages = _parse_docx(DocxDocument(path).paragraphs, splitter)
        for i, p in enumerate(passages):
            p.index = i
        return cls(passages, title=doc_title)

    @classmethod
    def _from_html(cls, path, splitter, title) -> 'Document':
        """Requires: beautifulsoup4"""
        from bs4 import BeautifulSoup
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        page_title = soup.title.string.strip() if soup.title and soup.title.string else None
        doc_title = title or page_title or os.path.splitext(os.path.basename(path))[0]
        passages = _parse_html(soup.body or soup, splitter)
        for i, p in enumerate(passages):
            p.index = i
        return cls(passages, title=doc_title)

    @classmethod
    def _from_epub(cls, path, splitter, title) -> 'Document':
        """Requires: ebooklib, beautifulsoup4"""
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        book = epub.read_epub(path, options={'ignore_ncx': True})
        doc_title = title or book.title or os.path.splitext(os.path.basename(path))[0]
        passages = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            passages.extend(_parse_html(soup.body or soup, splitter))
        for i, p in enumerate(passages):
            p.index = i
        return cls(passages, title=doc_title)

    @classmethod
    def _from_directory(cls, path, splitter, title) -> 'Document':
        doc_title = title or _path_to_title(path)
        logger.debug("Walking directory: %s", path)
        passages = []
        _walk_directory(path, passages, splitter, label_parts=[])
        for i, p in enumerate(passages):
            p.index = i
        doc = cls(passages, title=doc_title)
        logger.info("Loaded '%s' (directory): %d passages", doc_title, len(passages))
        return doc


# ---------------------------------------------------------------------------
# Structured format parsers (return list[Passage] with index=0, label set)
# ---------------------------------------------------------------------------

def _parse_markdown(lines: list[str], splitter: str | Callable) -> list[Passage]:
    """Parse a markdown file into passages, tracking heading hierarchy as labels."""
    passages: list[Passage] = []
    heading_stack: list[tuple[int, str]] = []  # (depth, title)
    pending_lines: list[str] = []

    def current_label() -> str | None:
        if not heading_stack:
            return None
        return " > ".join(title for _, title in heading_stack)

    def flush():
        label = current_label()
        for t in _apply_splitter(pending_lines, splitter):
            if t.strip():
                passages.append(Passage(text=t, index=0, label=label))
        pending_lines.clear()

    for line in lines:
        m = re.match(r'^(#+)\s+(.*)', line)
        if m:
            flush()
            depth = len(m.group(1))
            heading_title = m.group(2).strip()
            while heading_stack and heading_stack[-1][0] >= depth:
                heading_stack.pop()
            heading_stack.append((depth, heading_title))
        else:
            pending_lines.append(line)

    flush()
    return passages


def _parse_pdf(reader, splitter: str | Callable) -> list[Passage]:
    """Parse a PdfReader into passages, using page numbers as labels."""
    passages: list[Passage] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ''
        label = f"Page {i}"
        for t in _apply_splitter(text.splitlines(), splitter):
            if t.strip():
                passages.append(Passage(text=t.strip(), index=0, label=label))
    return passages


def _parse_docx(paragraphs, splitter: str | Callable) -> list[Passage]:
    """Parse python-docx paragraph objects into passages, tracking heading hierarchy."""
    passages: list[Passage] = []
    heading_stack: list[tuple[int, str]] = []
    pending_lines: list[str] = []

    def current_label() -> str | None:
        if not heading_stack:
            return None
        return " > ".join(text for _, text in heading_stack)

    def flush():
        label = current_label()
        for t in _apply_splitter(pending_lines, splitter):
            if t.strip():
                passages.append(Passage(text=t, index=0, label=label))
        pending_lines.clear()

    for para in paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue
        if style == 'Title':
            flush()
            heading_stack.clear()
            heading_stack.append((0, text))
        elif style.startswith('Heading '):
            flush()
            try:
                depth = int(style.split()[-1])
            except ValueError:
                depth = 1
            while heading_stack and heading_stack[-1][0] >= depth:
                heading_stack.pop()
            heading_stack.append((depth, text))
        else:
            pending_lines.append(text)

    flush()
    return passages


_HTML_HEADING_DEPTHS = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6}


def _parse_html(root, splitter: str | Callable) -> list[Passage]:
    """Parse a BeautifulSoup element into passages, tracking heading hierarchy."""
    passages: list[Passage] = []
    heading_stack: list[tuple[int, str]] = []
    pending_lines: list[str] = []

    def current_label() -> str | None:
        if not heading_stack:
            return None
        return " > ".join(text for _, text in heading_stack)

    def flush():
        label = current_label()
        for t in _apply_splitter(pending_lines, splitter):
            if t.strip():
                passages.append(Passage(text=t, index=0, label=label))
        pending_lines.clear()

    for element in root.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
        tag = element.name
        text = element.get_text(separator=' ', strip=True)
        if not text:
            continue
        if tag in _HTML_HEADING_DEPTHS:
            flush()
            depth = _HTML_HEADING_DEPTHS[tag]
            while heading_stack and heading_stack[-1][0] >= depth:
                heading_stack.pop()
            heading_stack.append((depth, text))
        else:
            pending_lines.append(text)

    flush()
    return passages


# ---------------------------------------------------------------------------
# Directory walker
# ---------------------------------------------------------------------------

# Extensions that require explicit schema params — skipped during directory walks.
_DATA_EXTENSIONS = {'.csv', '.tsv', '.json', '.jsonl', '.ndjson'}


def _walk_directory(
    dirpath: str,
    passages: list[Passage],
    splitter: str | Callable,
    label_parts: list[str],
):
    """Recursively walk a directory, collecting passages with path-based labels.

    Markdown, plain text, PDF, Word, HTML, and EPUB files are loaded automatically.
    CSV, TSV, JSON, and JSONL files are skipped — use from_csv()/from_json()/
    from_jsonl() directly for those formats.
    """
    for name in sorted(os.listdir(dirpath)):
        if name.startswith('.'):
            continue  # skip hidden files (.DS_Store, .gitignore, etc.)
        path = os.path.join(dirpath, name)
        file_title = _path_to_title(name)
        if os.path.isdir(path):
            _walk_directory(path, passages, splitter, label_parts + [file_title])
            continue
        if not os.path.isfile(path):
            continue

        ext = os.path.splitext(name)[1].lower()
        if ext in _DATA_EXTENSIONS:
            continue  # schema-dependent; use named constructors instead

        label_prefix = " > ".join(label_parts + [file_title])

        if ext == '.md':
            with open(path, 'r') as f:
                lines = [line.rstrip('\n') for line in f]
            for p in _parse_markdown(lines, splitter):
                p.label = f"{label_prefix} > {p.label}" if p.label else label_prefix
                passages.append(p)

        elif ext == '.pdf':
            from pypdf import PdfReader
            for p in _parse_pdf(PdfReader(path), splitter):
                p.label = f"{label_prefix} > {p.label}" if p.label else label_prefix
                passages.append(p)

        elif ext == '.docx':
            from docx import Document as DocxDocument
            for p in _parse_docx(DocxDocument(path).paragraphs, splitter):
                p.label = f"{label_prefix} > {p.label}" if p.label else label_prefix
                passages.append(p)

        elif ext in ('.html', '.htm'):
            from bs4 import BeautifulSoup
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            for p in _parse_html(soup.body or soup, splitter):
                p.label = f"{label_prefix} > {p.label}" if p.label else label_prefix
                passages.append(p)

        elif ext == '.epub':
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            book = epub.read_epub(path, options={'ignore_ncx': True})
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                for p in _parse_html(soup.body or soup, splitter):
                    p.label = f"{label_prefix} > {p.label}" if p.label else label_prefix
                    passages.append(p)

        else:
            try:
                with open(path, 'r') as f:
                    content = f.read()
                for t in _apply_splitter(content.splitlines(), splitter):
                    if t.strip():
                        passages.append(Passage(text=t, index=0, label=label_prefix))
            except UnicodeDecodeError:
                logger.warning("Skipping binary file (not UTF-8): %s", path)


# ---------------------------------------------------------------------------
# Splitter and path utilities
# ---------------------------------------------------------------------------

def _apply_splitter(lines: list[str], splitter: str | Callable) -> list[str]:
    """Split a list of lines into passage strings according to the splitter strategy."""
    if callable(splitter):
        result = splitter("\n".join(lines))
        return result if isinstance(result, list) else [result]

    if splitter == "line":
        return [line.strip() for line in lines]

    if splitter == "paragraph":
        paragraphs, current = [], []
        for line in lines:
            if line.strip():
                current.append(line.strip())
            elif current:
                paragraphs.append(" ".join(current))
                current = []
        if current:
            paragraphs.append(" ".join(current))
        return paragraphs

    if splitter == "section":
        # All content under the current heading becomes one passage
        joined = " ".join(line.strip() for line in lines if line.strip())
        return [joined] if joined else []

    if splitter == "sentence":
        import nltk
        nltk.download('punkt_tab', quiet=True)
        text = " ".join(line.strip() for line in lines if line.strip())
        return nltk.sent_tokenize(text)

    raise ValueError(
        f"Unknown splitter: {splitter!r}. "
        "Choose from: 'line', 'paragraph', 'section', 'sentence', or a callable."
    )


def _path_to_title(path: str) -> str:
    """Convert a file/directory path to a human-readable title.

    Strips numeric sort prefixes (e.g. '00_Section_1.txt' -> 'Section 1').
    """
    name = os.path.basename(path)
    name = os.path.splitext(name)[0]
    if name and name[0].isdigit():
        parts = name.split('_', 1)
        if len(parts) > 1:
            name = parts[1]
    return name.replace('_', ' ')
